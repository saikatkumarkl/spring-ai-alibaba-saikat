#!/usr/bin/env python3
"""
Generate comprehensive architecture documentation for Spring AI Alibaba Admin Platform.
Outputs: Spring_AI_Alibaba_Architecture.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    heading = doc.styles[f'Heading {level}']
    heading.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

# Code block style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Courier New'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(0.5)

# Caption style
caption_style = doc.styles.add_style('FigCaption', WD_STYLE_TYPE.PARAGRAPH)
caption_style.font.size = Pt(9)
caption_style.font.italic = True
caption_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code(text, lang="java"):
    """Add a code block to the document."""
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')


def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing


def add_note(text):
    """Add a highlighted note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"📌 Note: {text}")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x44, 0x00)


def add_flow_text(text):
    """Add a flow diagram as text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


# ═══════════════════════════════════════════════════════════════════════════
#                          TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Spring AI Alibaba Admin Platform', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Design, Architecture & Module-Call Documentation', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Technical Reference')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0 — June 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#                      TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary",
    "2. System Architecture Overview",
    "   2.1 High-Level Architecture",
    "   2.2 Module Structure",
    "   2.3 Technology Stack",
    "3. Docker Compose Infrastructure",
    "   3.1 Service Topology",
    "   3.2 Service Details",
    "   3.3 Network & Volume Configuration",
    "   3.4 Service Initialization Sequence",
    "4. Model & Provider Management",
    "   4.1 Provider Architecture",
    "   4.2 Model Auto-Sync Engine",
    "   4.3 ModelFactory — Universal Model Access",
    "   4.4 Database Schema",
    "   4.5 Frontend Model Configuration",
    "5. Knowledge Base System",
    "   5.1 Overview & Architecture",
    "   5.2 Knowledge Base CRUD",
    "   5.3 Document Upload & Processing Pipeline",
    "   5.4 Text Splitting & Chunking",
    "   5.5 Embedding & Vector Storage",
    "   5.6 Document Retrieval at Query Time",
    "   5.7 RAG Integration via KnowledgeBaseRetrievalAdvisor",
    "6. Source System (Document Crawling)",
    "   6.1 Overview & Architecture",
    "   6.2 Source System CRUD",
    "   6.3 ManifoldCF Bridge Service",
    "   6.4 CMIS Connector & ACL Extraction",
    "   6.5 REST API Connector",
    "   6.6 ManifoldCF Init Scripts",
    "   6.7 End-to-End Crawl Flow",
    "7. Chat System — How a User Message is Processed",
    "   7.1 Overview & Two Chat Paths",
    "   7.2 Request Entry Points (Controllers)",
    "   7.3 AgentServiceImpl — Orchestrator",
    "   7.4 BasicAgentExecutor — The Chat Engine",
    "   7.5 Building the ChatClient",
    "   7.6 Tool Execution",
    "   7.7 RAG Retrieval During Chat",
    "   7.8 Streaming Response",
    "   7.9 Complete Request Flow Diagram",
    "8. Agent Framework (spring-ai-alibaba-agent-framework)",
    "   8.1 Agent Hierarchy",
    "   8.2 ReactAgent — The Core Agent",
    "   8.3 AgentLlmNode — Model Invocation",
    "   8.4 AgentToolNode — Tool Execution",
    "   8.5 Multi-Agent Patterns",
    "   8.6 Hook System (Context Engineering)",
    "   8.7 Interceptor Chain",
    "9. Graph Core (spring-ai-alibaba-graph-core)",
    "   9.1 StateGraph — Workflow Definition",
    "   9.2 CompiledGraph — Execution Engine",
    "   9.3 OverAllState — State Container",
    "   9.4 Node & Edge Model",
    "   9.5 KeyStrategy — State Merge Policies",
    "   9.6 Checkpoint & Persistence",
    "   9.7 Compilation Pipeline",
    "10. Observability & Tracing",
    "   10.1 Micrometer Observation API",
    "   10.2 Graph-Level Observations",
    "   10.3 Metrics & Counters",
    "   10.4 Auto-Configuration",
    "11. Studio Module (Debug UI)",
    "12. Frontend Architecture",
    "   12.1 Technology Stack",
    "   12.2 Package Structure",
    "   12.3 Routing & Navigation",
    "   12.4 Key Pages",
    "13. Appendix — File Reference",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Number 2')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Spring AI Alibaba Admin Platform is a production-grade, full-stack AI application management system '
    'that provides a visual interface for building, deploying, and managing AI-powered agents. The platform '
    'integrates Large Language Models (LLMs), Retrieval-Augmented Generation (RAG), document crawling, '
    'and multi-agent orchestration into a unified, extensible architecture.'
)

doc.add_paragraph(
    'The system is built on Spring Boot 3.5.x with a React frontend (UmiJS 4 + Ant Design) and uses '
    'Docker Compose to orchestrate 13+ services including PostgreSQL, OpenSearch, Redis, RocketMQ, '
    'Ollama (local LLM inference), Nacos (service discovery), and Apache ManifoldCF (document crawling).'
)

doc.add_paragraph('Key capabilities include:')
items = [
    'Multi-provider model management — Unified access to Ollama, DashScope, OpenAI, DeepSeek via OpenAI-compatible API',
    'Knowledge Base system — Upload, parse, chunk, embed, and store documents for RAG retrieval',
    'Source System — External document crawling (CMIS, REST API) with ACL-aware indexing via ManifoldCF',
    'Chat with RAG — Streaming chat with automatic knowledge retrieval and tool calling',
    'Agent Framework — ReAct agents with hook-based context engineering and multi-agent orchestration',
    'Graph Core — Low-level workflow engine with state management, checkpointing, and parallel execution',
    'Observability — Micrometer-based tracing and metrics for graph execution',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 High-Level Architecture', level=2)

doc.add_paragraph(
    'The platform follows a layered architecture with clear separation between the presentation tier '
    '(React frontend), the application tier (Spring Boot backend), and the data/infrastructure tier '
    '(PostgreSQL, OpenSearch, Redis, RocketMQ, ManifoldCF, Ollama).'
)

add_code("""
┌─────────────────────────────────────────────────────────────────────────┐
│                        PRESENTATION TIER                                │
│  React 18 + UmiJS 4 + Ant Design + TypeScript                          │
│  ┌───────────┐  ┌──────────┐  ┌──────────┐  ┌────────┐  ┌──────────┐  │
│  │ Chat UI   │  │ Agent    │  │ Knowledge│  │ Source │  │ Settings │  │
│  │           │  │ Builder  │  │ Base     │  │ System │  │ (Models) │  │
│  └─────┬─────┘  └────┬─────┘  └────┬─────┘  └───┬────┘  └────┬─────┘  │
│        └──────────────┴──────────────┴────────────┴────────────┘        │
│                          HTTP/SSE (port 8000 → proxy → 8080)           │
├─────────────────────────────────────────────────────────────────────────┤
│                        APPLICATION TIER                                 │
│  Spring Boot 3.5.x (port 8080)                                         │
│  ┌──────────────────────────────────────────────────────────────────┐   │
│  │  Controllers: ChatController, KnowledgeBaseController,          │   │
│  │               SourceSystemController, ProviderController,       │   │
│  │               ModelController, AppController                    │   │
│  ├──────────────────────────────────────────────────────────────────┤   │
│  │  Services: AgentServiceImpl → BasicAgentExecutor                │   │
│  │            KnowledgeBaseServiceImpl, DocumentServiceImpl        │   │
│  │            SourceSystemServiceImpl, ManifoldCFBridgeService     │   │
│  │            ProviderManager, ModelManager, ModelFactory           │   │
│  ├──────────────────────────────────────────────────────────────────┤   │
│  │  Spring AI Integration:                                         │   │
│  │    ChatClient → ChatModel (OpenAI-compatible) → LLM Provider   │   │
│  │    EmbeddingModel → Vector embeddings                           │   │
│  │    VectorStore (OpenSearch) → Similarity search                 │   │
│  │    Advisor (RAG, Memory, Tool calling)                          │   │
│  └──────────────────────────────────────────────────────────────────┘   │
├─────────────────────────────────────────────────────────────────────────┤
│                    DATA / INFRASTRUCTURE TIER                           │
│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌─────────────┐  │
│  │PostgreSQL│ │OpenSearch │ │  Redis   │ │ RocketMQ │ │  ManifoldCF │  │
│  │  (5433)  │ │  (9200)  │ │  (6379)  │ │  (9876)  │ │   (8345)    │  │
│  │  RDBMS   │ │ Vectors  │ │  Cache   │ │  Async   │ │  Crawling   │  │
│  └──────────┘ └──────────┘ └──────────┘ └──────────┘ └─────────────┘  │
│  ┌──────────┐ ┌──────────┐ ┌──────────────────┐                       │
│  │  Ollama  │ │  Nacos   │ │ OpenSearch        │                       │
│  │ (11434)  │ │  (8848)  │ │ Dashboards (5601)│                       │
│  │ Local LLM│ │ Registry │ │ Search UI        │                       │
│  └──────────┘ └──────────┘ └──────────────────┘                       │
└─────────────────────────────────────────────────────────────────────────┘
""")

doc.add_heading('2.2 Module Structure', level=2)

doc.add_paragraph(
    'The project is organized as a Maven multi-module monorepo. The following table describes each module '
    'and its role in the system:'
)

add_table(
    ['Module', 'Purpose', 'Key Dependencies'],
    [
        ['spring-ai-alibaba-admin-server-start', 'Spring Boot application entry point, controllers, auto-config', 'server-core, server-runtime, server-openapi'],
        ['spring-ai-alibaba-admin-server-core', 'Domain models, DTOs, repository interfaces, utility classes', 'MyBatis-Plus, Spring Data'],
        ['spring-ai-alibaba-admin-server-runtime', 'Service implementations, business logic, pipeline processors', 'server-core, Spring AI, RocketMQ'],
        ['spring-ai-alibaba-admin-server-openapi', 'External API definitions, Swagger docs', 'server-core'],
        ['spring-ai-alibaba-agent-framework', 'High-level agent APIs: ReactAgent, SequentialAgent, ParallelAgent', 'graph-core, Spring AI'],
        ['spring-ai-alibaba-graph-core', 'Low-level workflow engine: StateGraph, CompiledGraph, Node, Edge', 'None (standalone)'],
        ['spring-ai-alibaba-studio', 'Embedded debug UI for visualizing agent execution', 'agent-framework, graph-core'],
        ['spring-boot-starters/*', 'Auto-configuration starters (Nacos, observation, built-in nodes)', 'graph-core, Micrometer'],
        ['spring-ai-alibaba-bom', 'Bill of Materials — centralized version management', 'N/A'],
        ['frontend/', 'React Admin UI (UmiJS 4, Ant Design, TypeScript)', 'Node.js >= 20'],
        ['manifoldcf-saikat/', 'Apache ManifoldCF fork with CMIS + REST API connectors', 'PostgreSQL, Jetty'],
    ]
)

doc.add_heading('2.3 Technology Stack', level=2)

add_table(
    ['Layer', 'Technology', 'Version', 'Purpose'],
    [
        ['Language', 'Java', '17', 'Backend development'],
        ['Framework', 'Spring Boot', '3.5.x', 'Application framework'],
        ['AI Framework', 'Spring AI', '1.0.x', 'LLM/embedding/vector abstraction'],
        ['Frontend', 'React + UmiJS + TypeScript', '18 / 4 / 5', 'Admin UI'],
        ['Database', 'PostgreSQL', '16', 'Relational data store'],
        ['Vector Store', 'OpenSearch', '3.4.0', 'KNN vector search + full-text search'],
        ['Cache', 'Redis', '7.x', 'Session/model/KB cache'],
        ['Message Queue', 'RocketMQ', '5.3.2', 'Async document indexing'],
        ['LLM Runtime', 'Ollama', 'Latest', 'Local model inference'],
        ['Service Registry', 'Nacos', '2.3.0', 'A2A service discovery'],
        ['Document Crawler', 'ManifoldCF', '2.30-SNAPSHOT', 'CMIS + REST API crawling'],
        ['ORM', 'MyBatis-Plus', '3.x', 'Database access'],
        ['Build', 'Maven', '3.x', 'Project build'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  3. DOCKER COMPOSE INFRASTRUCTURE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Docker Compose Infrastructure', level=1)

doc.add_paragraph(
    'The entire platform runs via Docker Compose defined in '
    'spring-ai-alibaba-admin/docker/middleware/docker-compose-arm.yaml. '
    'This file orchestrates 13+ services with health checks, dependency ordering, and named volumes.'
)

doc.add_heading('3.1 Service Topology', level=2)

add_code("""
                    ┌──────────────┐
                    │   Frontend   │ (localhost:8000)
                    │  React/UmiJS │
                    └──────┬───────┘
                           │ proxy
                    ┌──────▼───────┐
               ┌────┤   Backend    ├────┐
               │    │ Spring Boot  │    │
               │    │  (8080)      │    │
               │    └──┬──┬──┬──┬──┘    │
               │       │  │  │  │       │
    ┌──────────▼┐  ┌───▼┐ │  │ ┌▼─────┐ │  ┌───────────┐
    │PostgreSQL │  │Redis│ │  │ │Nacos │ │  │ RocketMQ  │
    │  (5433)   │  │6379 │ │  │ │8848  │ │  │NameSrv    │
    └───────────┘  └─────┘ │  │ └──────┘ │  │Broker     │
                           │  │          │  │Proxy      │
              ┌────────────▼┐ │          │  └───────────┘
              │  OpenSearch  │ │          │
              │   (9200)     │◄┘         │
              └──────┬───────┘           │
                     │                   │
              ┌──────▼───────┐    ┌──────▼───────┐
              │  ManifoldCF  │    │    Ollama     │
              │   (8345)     │    │   (11434)     │
              └──────────────┘    └───────────────┘
""")

doc.add_heading('3.2 Service Details', level=2)

services = [
    ['postgres', 'postgres:16', '5433:5432', 'Relational DB for admin app, ManifoldCF', 'pg_isready -U admin'],
    ['opensearch', 'opensearchproject/opensearch:3.4.0', '9200:9200, 9600:9600', 'Vector store + full-text search', 'curl http://localhost:9200/_cluster/health'],
    ['opensearch-dashboards', 'opensearchproject/opensearch-dashboards:3.4.0', '5601:5601', 'OpenSearch visualization UI', 'N/A'],
    ['nacos', 'nacos/nacos-server:v2.3.0', '8848, 9848, 9849', 'Service registry for A2A protocol', 'N/A'],
    ['redis', 'redis:7-alpine', '6379:6379', 'Cache (models, KBs, sessions)', 'redis-cli ping'],
    ['rmq-namesrv', 'apache/rocketmq:5.3.2', '9876:9876', 'RocketMQ name server', 'ps aux | grep mqnamesrv'],
    ['rmq-broker', 'apache/rocketmq:5.3.2', '10909-10912', 'RocketMQ message broker', 'ps aux | grep mqbroker'],
    ['rmq-proxy', 'apache/rocketmq:5.3.2', '18080-18081', 'RocketMQ gRPC proxy', 'N/A'],
    ['rmq-init-topic', 'apache/rocketmq:5.3.2', 'N/A', 'Creates topic: topic_saa_studio_document_index', 'N/A (one-shot)'],
    ['ollama', 'ollama/ollama', '11434:11434', 'Local LLM inference server', 'ollama list'],
    ['ollama-init', 'ollama/ollama', 'N/A', 'Pulls models from models.conf', 'N/A (one-shot)'],
    ['backend', 'Built from Dockerfile.backend', '8080:8080', 'Spring Boot admin server', 'wget actuator/health'],
    ['manifoldcf', 'Built from Dockerfile.maven', '8345:8345', 'Document crawler (CMIS + REST API)', 'curl /mcf-api-service/json/outputconnectors'],
    ['manifoldcf-init', 'curlimages/curl:8.5.0', 'N/A', '6-step MCF auto-configuration', 'N/A (one-shot)'],
]

add_table(['Service', 'Image', 'Ports', 'Purpose', 'Health Check'], services)

doc.add_heading('3.3 Network & Volume Configuration', level=2)

doc.add_paragraph('All services connect to a single Docker bridge network: saa-network')
doc.add_paragraph('Named volumes for data persistence:')
items = [
    'opensearch_data — OpenSearch indices and vector data',
    'saa_storage — Backend file storage (uploaded documents)',
    'ollama/data — Ollama model weights',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.4 Service Initialization Sequence', level=2)

doc.add_paragraph(
    'Services start in a strict dependency order enforced by depends_on with health checks:'
)

add_code("""
Phase 1 (Infrastructure):  postgres, opensearch, redis → start independently
Phase 2 (Messaging):       rmq-namesrv → rmq-broker → rmq-proxy → rmq-init-topic
Phase 3 (LLM):             ollama → ollama-init (pulls models from models.conf)
Phase 4 (Service Registry): nacos (standalone mode)
Phase 5 (Application):     backend (waits for postgres, redis, opensearch, rmq-proxy, ollama)
Phase 6 (Crawling):        manifoldcf (waits for postgres, opensearch)
Phase 7 (Configuration):   manifoldcf-init (waits for manifoldcf healthy)
                            → Step 0: Create OpenSearch index template (manifoldcf-acl)
                            → Step 1: Create OpenSearch output connection
                            → Step 2: Create CMIS repository connection
                            → Step 3: Create authorities sync job
                            → Step 4: Create document crawl job
                            → Step 5: Start the crawl job
""")

doc.add_heading('3.4.1 Backend Environment Variables', level=3)

doc.add_paragraph('The backend service receives its configuration via environment variables:')

add_code("""
# docker-compose-arm.yaml → backend service
environment:
  SPRING_PROFILES_ACTIVE: docker
  JAVA_OPTS: "-Xms512m -Xmx1024m -XX:+UseG1GC"
  SPRING_DATASOURCE_URL: jdbc:postgresql://postgres:5432/admin
  SPRING_DATASOURCE_USERNAME: admin
  SPRING_DATASOURCE_PASSWORD: admin
  SPRING_REDIS_HOST: redis
  SPRING_REDIS_PORT: 6379
  SPRING_ELASTICSEARCH_URL: http://opensearch:9200
  NACOS_SERVER_ADDR: nacos:8848
  ROCKETMQ_ENDPOINTS: rmq-proxy:18081
  OLLAMA_BASE_URL: http://ollama:11434
  OLLAMA_CHAT_MODEL: ${OLLAMA_CHAT_MODEL:-qwen2.5:7b}
  MODEL_CONFIG_FILE: /app/model-config.yml
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  4. MODEL & PROVIDER MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Model & Provider Management', level=1)

doc.add_paragraph(
    'The platform manages AI model providers (Ollama, DashScope, OpenAI, DeepSeek, etc.) and their '
    'models through a unified abstraction. All providers are accessed via the OpenAI-compatible API, '
    'enabling a single code path for all LLM interactions.'
)

doc.add_heading('4.1 Provider Architecture', level=2)

add_code("""
Frontend (Settings/ModelService)          Backend
┌──────────────────────────┐              ┌──────────────────────────────────┐
│ ProviderCard             │──GET ────────│ ProviderController               │
│ ProviderInfoForm         │──PUT ────────│   /console/v1/providers          │
│ ModelConfigModal         │──POST ───────│                                  │
│ ModelSelector (dropdown) │──GET ────────│ ModelController                  │
│                          │              │   /console/v1/models             │
└──────────────────────────┘              └──────────┬───────────────────────┘
                                                     │
                                          ┌──────────▼───────────────────────┐
                                          │ ProviderManager / ModelManager   │
                                          │ (MyBatis-Plus + Redis Cache)     │
                                          └──────────┬───────────────────────┘
                                                     │
                                          ┌──────────▼───────────────────────┐
                                          │ PostgreSQL                       │
                                          │ Tables: provider, model          │
                                          └──────────────────────────────────┘
""")

doc.add_heading('4.1.1 Provider Entity (Database Schema)', level=3)

add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'Long (AUTO)', 'Primary key'],
        ['workspace_id', 'String', 'Tenant isolation'],
        ['icon', 'String', 'Provider icon URL'],
        ['name', 'String', 'Display name (e.g., "Ollama")'],
        ['code', 'String', 'Unique 8-char UUID code'],
        ['enable', 'Integer', '0=disabled, 1=enabled'],
        ['protocol', 'String', 'Default "openai"'],
        ['source', 'String', '"custom" or "preset"'],
        ['supported_model_types', 'String', 'Comma-separated: llm, embedding, rerank, tts, stt'],
        ['credential', 'String', 'JSON: apiKey (RSA encrypted), endpoint'],
    ]
)

doc.add_heading('4.1.2 Model Entity (Database Schema)', level=3)

add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'Long (AUTO)', 'Primary key'],
        ['workspace_id', 'String', 'Tenant isolation'],
        ['name', 'String', 'Display name'],
        ['model_id', 'String', 'Unique identifier (e.g., llama3.2:3b)'],
        ['provider_code', 'String', 'References provider.code'],
        ['type', 'String', 'llm, embedding, rerank, tts, stt'],
        ['enable', 'Integer', '0=disabled, 1=enabled'],
        ['tags', 'String', 'Comma-separated: general, vision, coding, reasoning'],
        ['source', 'String', '"custom", "preset", or "auto"'],
    ]
)

doc.add_heading('4.2 Model Auto-Sync Engine', level=2)

doc.add_paragraph(
    'When the Model Selector API is called (GET /console/v1/models/{type}/selector), the system '
    'automatically syncs models from each provider\'s live endpoint. This ensures the dropdown '
    'always reflects the actual models available.'
)

add_code("""
// ModelController.java — getModelSelector()
@GetMapping("/{type}/selector")
public Result<Map<String, List<ModelConfigInfo>>> getModelSelector(@PathVariable String type) {
    // 1. Query all providers
    List<ProviderConfigInfo> providers = providerManager.queryProviders(workspaceId, null);
    
    // 2. For each ENABLED provider, auto-sync
    for (ProviderConfigInfo provider : providers) {
        if (provider.getEnable()) {
            autoSyncModelsFromProvider(provider);
        }
    }
    
    // 3. Query all models, filter by type, group by provider
    List<ModelConfigInfo> models = modelManager.queryModels(workspaceId, type);
    return Result.success(groupByProvider(models));
}
""")

doc.add_paragraph('The auto-sync flow works as follows:')

add_code("""
// Auto-sync: fetchLiveModels()
private List<String> fetchLiveModels(ProviderConfigInfo provider) {
    String endpoint = provider.getCredential().getEndpoint();
    
    // Try 1: OpenAI-compatible /v1/models
    try {
        ResponseEntity<Map> response = restTemplate.getForEntity(
            endpoint + "/v1/models", Map.class);
        return parseOpenAiModelList(response.getBody());
    } catch (Exception e) {
        // Try 2: Ollama-native /api/tags
        ResponseEntity<Map> response = restTemplate.getForEntity(
            endpoint + "/api/tags", Map.class);
        return parseOllamaModelList(response.getBody());
    }
}

// Auto-sync: inferModelType()
private String inferModelType(String modelName) {
    String lower = modelName.toLowerCase();
    if (lower.contains("embed") || lower.contains("bge-") 
        || lower.contains("all-minilm") || lower.contains("e5-")) {
        return "embedding";
    }
    if (lower.contains("rerank")) return "rerank";
    return "llm";
}
""")

doc.add_heading('4.3 ModelFactory — Universal Model Access', level=2)

doc.add_paragraph(
    'ALL model providers are accessed through a single factory that creates OpenAI-compatible '
    'client instances. This is the key architectural decision that enables multi-provider support '
    'without provider-specific code.'
)

add_code("""
// ModelFactory.java — getChatModel()
public ChatModel getChatModel(String providerCode, String modelId) {
    String cacheKey = providerCode + ":" + modelId;
    return chatModelCache.computeIfAbsent(cacheKey, k -> {
        ProviderConfigInfo provider = providerManager.getProvider(providerCode);
        ModelCredential credential = provider.getCredential();
        
        // CRITICAL: ALL providers use OpenAI-compatible API
        OpenAiApi api = OpenAiApi.builder()
            .baseUrl(credential.getEndpoint())   // e.g., http://ollama:11434/v1
            .apiKey(credential.getApiKey())       // or "ollama" for local
            .build();
            
        return OpenAiChatModel.builder()
            .openAiApi(api)
            .defaultOptions(ChatOptionsBuilder.builder()
                .model(modelId)
                .build())
            .build();
    });
}

// ModelFactory.java — getEmbeddingModel()
public EmbeddingModel getEmbeddingModel(String providerCode, String modelId) {
    String cacheKey = providerCode + ":" + modelId;
    return embeddingModelCache.computeIfAbsent(cacheKey, k -> {
        ProviderConfigInfo provider = providerManager.getProvider(providerCode);
        ModelCredential credential = provider.getCredential();
        
        OpenAiApi api = OpenAiApi.builder()
            .baseUrl(credential.getEndpoint())
            .apiKey(credential.getApiKey())
            .build();
            
        return new OpenAiEmbeddingModel(api,
            OpenAiEmbeddingOptions.builder()
                .model(modelId)
                .build());
    });
}
""")

add_note(
    'This design means Ollama models are accessed via http://ollama:11434/v1 — the /v1 suffix '
    'makes Ollama\'s API compatible with the OpenAI client. DashScope, DeepSeek, and other providers '
    'similarly expose OpenAI-compatible endpoints.'
)

doc.add_heading('4.4 Frontend Model Configuration', level=2)

doc.add_paragraph(
    'The frontend\'s Settings → Model Service page allows administrators to:'
)
items = [
    'Add/edit/delete model providers with API keys and endpoints',
    'Enable/disable providers and individual models',
    'View auto-synced models with inferred types and tags',
    'Manually add models with custom types and capabilities',
    'Model Selector dropdowns throughout the app pull from this configuration',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  5. KNOWLEDGE BASE SYSTEM
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Knowledge Base System', level=1)

doc.add_heading('5.1 Overview & Architecture', level=2)

doc.add_paragraph(
    'The Knowledge Base system provides RAG (Retrieval-Augmented Generation) capabilities. Users '
    'create knowledge bases, upload documents, which are then parsed, chunked, embedded, and stored '
    'in OpenSearch for vector similarity search during chat.'
)

add_code("""
 Upload Flow:                                       Retrieval Flow:
 ┌────────┐    ┌────────────┐    ┌──────────────┐  ┌──────────┐    ┌───────────┐
 │Frontend │───►│ Document   │───►│  RocketMQ    │  │Chat Req  │───►│Retrieval  │
 │ Upload  │    │ Controller │    │  Message     │  │          │    │ Advisor   │
 └────────┘    └────────────┘    └──────┬───────┘  └──────────┘    └─────┬─────┘
                                        │                                │
                                 ┌──────▼───────┐              ┌────────▼────────┐
                                 │ Document      │              │ KB Document     │
                                 │ IndexHandler  │              │ Retriever       │
                                 │ (MQ Consumer) │              │ (per KB, ∥)     │
                                 └──────┬───────┘              └────────┬────────┘
                                        │                                │
                               ┌────────▼────────┐             ┌────────▼────────┐
                               │ KB Index        │             │  OpenSearch      │
                               │ Pipeline        │             │  VectorStore    │
                               │ parse→chunk→    │             │  (KNN search)   │
                               │ embed→store     │             └─────────────────┘
                               └────────┬────────┘
                                        │
                               ┌────────▼────────┐
                               │   OpenSearch     │
                               │  Vector Index    │
                               │  (knn_vector)    │
                               └─────────────────┘
""")

doc.add_heading('5.2 Knowledge Base CRUD', level=2)

doc.add_paragraph('REST API endpoints at /console/v1/knowledge-bases:')

add_table(
    ['Method', 'Path', 'Description'],
    [
        ['POST', '/', 'Create knowledge base (generates kbId, creates OpenSearch index)'],
        ['PUT', '/{kbId}', 'Update knowledge base metadata'],
        ['DELETE', '/{kbId}', 'Delete KB (drops OpenSearch index, soft-delete)'],
        ['GET', '/{kbId}', 'Get knowledge base details'],
        ['GET', '/', 'List knowledge bases with pagination'],
        ['POST', '/{kbId}/upload', 'Upload documents to knowledge base'],
        ['POST', '/{kbId}/retrieve', 'Test retrieval against knowledge base'],
    ]
)

doc.add_heading('5.2.1 Creating a Knowledge Base', level=3)

add_code("""
// KnowledgeBaseServiceImpl.java — createKnowledgeBase()
public KnowledgeBase createKnowledgeBase(KnowledgeBase dto) {
    // 1. Generate unique kbId
    String kbId = IdGenerator.generateShortId();
    
    // 2. Set defaults for indexing config
    IndexConfig indexConfig = dto.getIndexConfig();
    // indexConfig contains: embeddingProvider, embeddingModel
    
    // 3. Set defaults for processing config
    ProcessConfig processConfig = dto.getProcessConfig();
    // processConfig contains: chunkType (STRUCTURE_AWARE | REGEX),
    //                         chunkSize (default 500), chunkOverlap (default 50)
    
    // 4. Create OpenSearch vector index
    int dimension = EmbeddingModelDimension.getDimension(
        indexConfig.getEmbeddingProvider(), 
        indexConfig.getEmbeddingModel()
    );
    openSearchVectorStoreService.createIndex(kbId, dimension);
    
    // 5. Save to PostgreSQL
    KnowledgeBaseEntity entity = convertToEntity(dto);
    knowledgeBaseMapper.insert(entity);
    
    // 6. Cache in Redis
    redisTemplate.opsForValue().set(
        KB_CACHE_KEY + kbId, entity, 12, TimeUnit.HOURS);
    
    return convertToDto(entity);
}
""")

doc.add_heading('5.2.2 OpenSearch Index Creation', level=3)

add_code("""
// OpenSearchVectorStoreService.java — createIndex()
public void createIndex(String kbId, int dimension) {
    String indexName = "kb_" + kbId;
    
    // KNN vector mapping with HNSW algorithm
    Map<String, Object> mapping = Map.of(
        "properties", Map.of(
            "embedding", Map.of(
                "type", "knn_vector",
                "dimension", dimension,
                "method", Map.of(
                    "name", "hnsw",
                    "space_type", "cosinesimil",
                    "engine", "lucene",
                    "parameters", Map.of(
                        "ef_construction", 128,
                        "m", 16
                    )
                )
            ),
            "content", Map.of("type", "text"),
            "metadata", Map.of("type", "object")
        )
    );
    
    CreateIndexRequest request = new CreateIndexRequest(indexName);
    request.mapping(mapping);
    openSearchClient.indices().create(request);
}
""")

doc.add_heading('5.3 Document Upload & Processing Pipeline', level=2)

doc.add_paragraph(
    'When a user uploads documents to a knowledge base, the processing happens asynchronously '
    'via RocketMQ to avoid blocking the HTTP request.'
)

add_code("""
// DocumentServiceImpl.java — createDocuments()
public List<Document> createDocuments(String kbId, List<MultipartFile> files) {
    List<Document> documents = new ArrayList<>();
    
    for (MultipartFile file : files) {
        // 1. Save file to storage
        String filePath = fileStorage.save(file);
        
        // 2. Create document record in DB
        DocumentEntity entity = new DocumentEntity();
        entity.setKbId(kbId);
        entity.setFileName(file.getOriginalFilename());
        entity.setFilePath(filePath);
        entity.setIndexStatus("PENDING");
        documentMapper.insert(entity);
        
        // 3. Send async message to RocketMQ
        DocumentIndexMessage message = new DocumentIndexMessage();
        message.setDocumentId(entity.getDocumentId());
        message.setKbId(kbId);
        rocketMQTemplate.sendAndForget(
            "topic_saa_studio_document_index", 
            message
        );
        
        documents.add(convertToDto(entity));
    }
    return documents;
}
""")

doc.add_heading('5.3.1 RocketMQ Consumer — DocumentIndexHandler', level=3)

add_code("""
// DocumentIndexHandler.java — handles async document indexing
@RocketMQMessageListener(
    topic = "topic_saa_studio_document_index",
    consumerGroup = "group_saa_studio_document_index"
)
public class DocumentIndexHandler implements RocketMQListener<DocumentIndexMessage> {
    
    @Override
    public void onMessage(DocumentIndexMessage message) {
        String docId = message.getDocumentId();
        String kbId = message.getKbId();
        
        try {
            // Update status to INDEXING
            documentMapper.updateIndexStatus(docId, "INDEXING");
            
            // Run the pipeline: parse → chunk → embed → store
            knowledgeBaseIndexPipeline.process(kbId, docId);
            
            // Update status to COMPLETED
            documentMapper.updateIndexStatus(docId, "COMPLETED");
            documentMapper.updateChunkCount(docId, 
                knowledgeBaseIndexPipeline.getLastChunkCount());
        } catch (Exception e) {
            documentMapper.updateIndexStatus(docId, "FAILED");
            documentMapper.updateErrorMessage(docId, e.getMessage());
            log.error("Failed to index document: {}", docId, e);
        }
    }
}
""")

doc.add_heading('5.4 Text Splitting & Chunking', level=2)

doc.add_paragraph(
    'The KnowledgeBaseIndexPipeline orchestrates the parse → transform → store stages. '
    'Two splitter strategies are available:'
)

add_table(
    ['Strategy', 'Class', 'Description'],
    [
        ['STRUCTURE_AWARE', 'StructureAwareTextSplitter (468 lines)', 'Preserves document structure (headings, paragraphs, lists). Splits at natural boundaries.'],
        ['REGEX', 'RegexTextSplitter', 'Splits by regex pattern (default: paragraph breaks). Simpler but less context-aware.'],
    ]
)

add_code("""
// KnowledgeBaseIndexPipeline.java — process()
public void process(String kbId, String documentId) {
    KnowledgeBase kb = knowledgeBaseService.getKnowledgeBase(kbId);
    DocumentEntity doc = documentMapper.selectByDocumentId(documentId);
    
    // 1. PARSE — Read file into Spring AI Document objects
    List<org.springframework.ai.document.Document> documents = parseDocument(doc);
    
    // 2. TRANSFORM — Split into chunks
    ProcessConfig config = kb.getProcessConfig();
    TextSplitter splitter;
    if ("STRUCTURE_AWARE".equals(config.getChunkType())) {
        splitter = new StructureAwareTextSplitter(
            config.getChunkSize(),      // default 500
            config.getChunkOverlap()    // default 50
        );
    } else {
        splitter = new RegexTextSplitter(
            config.getChunkSize(), 
            config.getChunkOverlap()
        );
    }
    List<org.springframework.ai.document.Document> chunks = splitter.apply(documents);
    
    // 3. STORE — Embed and persist to OpenSearch
    // OpenSearchVectorStore.add() internally calls EmbeddingModel to vectorize
    String indexName = "kb_" + kbId;
    OpenSearchVectorStore vectorStore = openSearchVectorStoreService
        .getVectorStore(kbId, kb.getIndexConfig());
    vectorStore.add(chunks);
    
    this.lastChunkCount = chunks.size();
}

private List<org.springframework.ai.document.Document> parseDocument(DocumentEntity doc) {
    String fileName = doc.getFileName().toLowerCase();
    DocumentReader reader;
    
    if (fileName.endsWith(".pdf")) {
        reader = new TikaDocumentReader(new FileSystemResource(doc.getFilePath()));
    } else if (fileName.endsWith(".md") || fileName.endsWith(".markdown")) {
        reader = new MarkdownDocumentReader(new FileSystemResource(doc.getFilePath()));
    } else {
        reader = new TextReader(new FileSystemResource(doc.getFilePath()));
    }
    return reader.get();
}
""")

doc.add_heading('5.5 Embedding & Vector Storage', level=2)

doc.add_paragraph(
    'When vectorStore.add(chunks) is called, Spring AI\'s OpenSearchVectorStore internally '
    'calls the configured EmbeddingModel to convert text chunks into vector embeddings, then '
    'stores them in OpenSearch\'s KNN index.'
)

add_code("""
// OpenSearchVectorStoreService.java — getVectorStore()
public OpenSearchVectorStore getVectorStore(String kbId, IndexConfig indexConfig) {
    String cacheKey = kbId;
    return vectorStoreCache.computeIfAbsent(cacheKey, k -> {
        // Get embedding model from ModelFactory
        EmbeddingModel embeddingModel = modelFactory.getEmbeddingModel(
            indexConfig.getEmbeddingProvider(),
            indexConfig.getEmbeddingModel()
        );
        
        String indexName = "kb_" + kbId;
        return OpenSearchVectorStore.builder(openSearchClient, embeddingModel)
            .index(indexName)
            .build();
    });
}
""")

doc.add_paragraph('The embedding dimension is determined by a lookup table:')

add_code("""
// EmbeddingModelDimension.java
public class EmbeddingModelDimension {
    private static final Map<String, Integer> DIMENSIONS = Map.of(
        "text-embedding-ada-002", 1536,
        "text-embedding-3-small", 1536,
        "text-embedding-3-large", 3072,
        "nomic-embed-text", 768,
        "all-minilm", 384,
        "bge-m3", 1024,
        "mxbai-embed-large", 1024
        // ... more models
    );
    
    public static int getDimension(String provider, String model) {
        return DIMENSIONS.getOrDefault(model, 1536); // default 1536
    }
}
""")

doc.add_heading('5.6 Document Retrieval at Query Time', level=2)

doc.add_paragraph(
    'During chat, the DocumentRetrieverManager creates per-knowledge-base retrievers that '
    'execute in parallel with a 30-second timeout.'
)

add_code("""
// DocumentRetrieverManager.java
public DocumentRetriever getDocumentRetriever(FileSearchOptions options) {
    List<String> kbIds = options.getKbIds();
    int topK = options.getTopK();                    // default 5
    double threshold = options.getSimilarityThreshold(); // default 0.5
    
    List<KnowledgeBaseDocumentRetriever> retrievers = kbIds.stream()
        .map(kbId -> new KnowledgeBaseDocumentRetriever(
            openSearchVectorStoreService.getVectorStore(kbId, getIndexConfig(kbId)),
            topK,
            threshold,
            options.getSearchType()  // SIMILARITY, HYBRID, FULL_TEXT
        ))
        .toList();
    
    return query -> {
        // Execute all retrievers in parallel with 30s timeout
        ExecutorService executor = Executors.newFixedThreadPool(retrievers.size());
        List<Future<List<Document>>> futures = retrievers.stream()
            .map(r -> executor.submit(() -> r.retrieve(query)))
            .toList();
        
        List<Document> allDocs = new ArrayList<>();
        for (Future<List<Document>> future : futures) {
            allDocs.addAll(future.get(30, TimeUnit.SECONDS));
        }
        
        // Optional: rerank results
        if (options.getEnableRerank()) {
            allDocs = rerankDocuments(allDocs, query, options);
        }
        
        // Filter by score threshold and return top-K
        return allDocs.stream()
            .filter(d -> d.getScore() >= threshold)
            .sorted(Comparator.comparingDouble(Document::getScore).reversed())
            .limit(topK)
            .toList();
    };
}
""")

doc.add_heading('5.7 RAG Integration via KnowledgeBaseRetrievalAdvisor', level=2)

doc.add_paragraph(
    'The KnowledgeBaseRetrievalAdvisor is a Spring AI Advisor that intercepts chat requests, '
    'retrieves relevant documents, and injects them into the system prompt before the LLM call.'
)

add_code("""
// KnowledgeBaseRetrievalAdvisor.java — adviseRequest()
@Override
public AdvisedRequest adviseRequest(AdvisedRequest request, Map<String, Object> context) {
    // 1. Extract user query from the latest message
    String userQuery = request.userText();
    
    // 2. Retrieve relevant documents
    List<Document> documents = documentRetriever.retrieve(userQuery);
    
    // 3. Format documents into context string
    String documentContext = documents.stream()
        .map(doc -> String.format("Source: %s\\nContent: %s",
            doc.getMetadata().get("source"), doc.getContent()))
        .collect(Collectors.joining("\\n\\n---\\n\\n"));
    
    // 4. Augment system prompt with retrieved documents
    // The system prompt template contains a {documents} placeholder
    String augmentedSystemPrompt = request.systemText()
        .replace("{documents}", documentContext);
    
    // 5. Store retrieval metadata for response
    context.put("retrieved_documents", documents);
    context.put("kb_ids", fileSearchOptions.getKbIds());
    
    return AdvisedRequest.from(request)
        .systemText(augmentedSystemPrompt)
        .build();
}
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  6. SOURCE SYSTEM (DOCUMENT CRAWLING)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Source System (Document Crawling)', level=1)

doc.add_heading('6.1 Overview & Architecture', level=2)

doc.add_paragraph(
    'The Source System feature allows connecting external document repositories (CMIS servers like '
    'Alfresco, SharePoint, or REST APIs like Confluence, Jira) for automated document crawling and '
    'indexing into OpenSearch. It bridges the Spring Boot admin backend with Apache ManifoldCF.'
)

add_code("""
 ┌──────────────┐        ┌────────────────────┐        ┌──────────────┐
 │   Frontend   │───────►│  SourceSystem      │───────►│ ManifoldCF   │
 │  Source UI   │  REST  │  Controller        │  REST  │ Bridge       │
 │              │        │  (11 endpoints)    │  API   │ Service      │
 └──────────────┘        └────────┬───────────┘        └──────┬───────┘
                                  │                           │
                         ┌────────▼───────────┐        ┌──────▼───────┐
                         │  SourceSystem      │        │ ManifoldCF   │
                         │  ServiceImpl       │        │ REST API     │
                         │  (CRUD + sync)     │        │ (port 8345)  │
                         └────────┬───────────┘        └──────┬───────┘
                                  │                           │
                         ┌────────▼───────────┐        ┌──────▼───────┐
                         │  PostgreSQL        │        │  OpenSearch   │
                         │  source_system     │        │  manifoldcf  │
                         │  table             │        │  index       │
                         └────────────────────┘        └──────────────┘
""")

doc.add_heading('6.2 Source System CRUD', level=2)

doc.add_paragraph('REST API endpoints at /console/v1/source-systems:')

add_table(
    ['Method', 'Path', 'Description'],
    [
        ['POST', '/', 'Create source system + MCF repository connection'],
        ['PUT', '/{sourceId}', 'Update source system configuration'],
        ['DELETE', '/{sourceId}', 'Delete source + abort MCF job + delete MCF connection'],
        ['GET', '/{sourceId}', 'Get source details (refreshes live MCF status)'],
        ['GET', '/', 'List sources with pagination'],
        ['POST', '/{sourceId}/sync', 'Start document sync (creates MCF crawl job)'],
        ['POST', '/{sourceId}/stop', 'Stop running sync job'],
        ['GET', '/{sourceId}/status', 'Get sync status from MCF'],
        ['POST', '/{sourceId}/test-connection', 'Test MCF connection (4-step validation)'],
        ['GET', '/connector-types', 'List available connector types (CMIS, REST API)'],
        ['GET', '/{sourceId}/documents', 'List crawled documents from OpenSearch'],
    ]
)

doc.add_heading('6.2.1 Creating a Source System', level=3)

add_code("""
// SourceSystemServiceImpl.java — createSourceSystem()
public SourceSystem createSourceSystem(SourceSystem dto) {
    // 1. Validate name uniqueness
    validateNameUniqueness(dto.getName(), dto.getWorkspaceId());
    
    // 2. Generate unique sourceId
    String sourceId = IdGenerator.generateShortId();
    dto.setSourceId(sourceId);
    
    // 3. Create ManifoldCF repository connection
    String mcfConnectionName = "source_" + sourceId;
    manifoldCFBridgeService.createRepositoryConnection(
        mcfConnectionName,
        dto.getConnectorType(),    // "cmis" or "restapi"
        dto.getConnectorConfig()   // CMIS/REST API specific parameters
    );
    dto.setMcfConnectionName(mcfConnectionName);
    
    // 4. Save to PostgreSQL
    SourceSystemEntity entity = convertToEntity(dto);
    entity.setSyncStatus("IDLE");
    sourceSystemMapper.insert(entity);
    
    return convertToDto(entity);
}
""")

doc.add_heading('6.3 ManifoldCF Bridge Service', level=2)

doc.add_paragraph(
    'ManifoldCFBridgeService is an HTTP client that wraps ManifoldCF\'s REST API using '
    'Spring\'s RestTemplate. It translates admin platform operations into MCF API calls.'
)

add_code("""
// ManifoldCFBridgeService.java

@Service
public class ManifoldCFBridgeService {
    private final RestTemplate restTemplate;
    private final String mcfApiUrl;  // http://manifoldcf:8345/mcf-api-service/json
    
    // Creates a repository connection in ManifoldCF
    public void createRepositoryConnection(String name, String type, 
                                           Map<String, Object> config) {
        // MCF uses a specific JSON format with _children_ arrays
        Map<String, Object> payload = Map.of(
            "repositoryconnection", Map.of(
                "_children_", List.of(
                    Map.of("_type_", "name", "_value_", name),
                    Map.of("_type_", "class_name", "_value_", 
                        getConnectorClassName(type)),
                    Map.of("_type_", "description", "_value_", 
                        "Auto-created by SAA"),
                    Map.of("_type_", "max_connections", "_value_", "10"),
                    Map.of("_type_", "configuration", 
                        "_children_", buildConfigParams(type, config))
                )
            )
        );
        
        restTemplate.put(
            mcfApiUrl + "/repositoryconnections/" + name, 
            payload
        );
    }
    
    // Creates a crawl job referencing the repository connection
    public String createCrawlJob(String connectionName, String outputConnection,
                                 Map<String, Object> jobConfig) {
        Map<String, Object> payload = buildJobPayload(
            connectionName, outputConnection, jobConfig);
        
        ResponseEntity<Map> response = restTemplate.postForEntity(
            mcfApiUrl + "/jobs", payload, Map.class);
        
        return extractJobId(response.getBody());
    }
    
    // Starts a crawl job
    public void startJob(String jobId) {
        restTemplate.put(
            mcfApiUrl + "/start/" + jobId, null);
    }
    
    // Gets live job status from ManifoldCF
    public Map<String, Object> getJobStatus(String jobId) {
        ResponseEntity<Map> response = restTemplate.getForEntity(
            mcfApiUrl + "/jobstatuses/" + jobId, Map.class);
        return response.getBody();
    }
}
""")

doc.add_heading('6.4 CMIS Connector & ACL Extraction', level=2)

doc.add_paragraph(
    'The CMIS connector in ManifoldCF connects to CMIS-compliant repositories (Alfresco, SharePoint, '
    'IBM FileNet, Nuxeo) and extracts documents with their ACLs (Access Control Lists).'
)

add_code("""
// CmisRepositoryConnector.java — extractAndSetAcl()
// Located in: manifoldcf-saikat/connectors/cmis/connector/

private void extractAndSetAcl(RepositoryDocument rd, Session session,
                               CmisObject cmisObject) {
    // CRITICAL: Cannot use cmisObject.getAcl() — returns null because
    // default OperationContext has includeACL=false.
    // Must use binding-level API for explicit ACL retrieval.
    
    Acl acl = session.getBinding().getAclService().getAcl(
        session.getRepositoryInfo().getId(),
        cmisObject.getId(),
        true,    // onlyBasicPermissions
        null     // extension
    );
    
    if (acl != null) {
        List<String> allowTokens = new ArrayList<>();
        List<String> denyTokens = new ArrayList<>();
        
        for (Ace ace : acl.getAces()) {
            String principal = ace.getPrincipalId().toLowerCase();
            
            if (ace.getPermissions().stream()
                    .anyMatch(p -> p.contains("cmis:read") || p.contains("Read"))) {
                if (ace.isDirect()) {
                    allowTokens.add(principal);
                }
            }
        }
        
        // Set document-level ACLs
        rd.setSecurity(RepositoryDocument.SECURITY_TYPE_DOCUMENT,
            allowTokens.toArray(new String[0]),
            denyTokens.toArray(new String[0])
        );
    }
    
    // Also extract parent folder ACLs
    extractParentFolderAcl(rd, session, cmisObject);
}
""")

add_note(
    'ACL tokens are stored raw in OpenSearch as allow_token_document and deny_token_document fields. '
    'At query time, the search filter includes the user\'s identity and group memberships to enforce '
    'document-level security.'
)

doc.add_heading('6.5 REST API Connector', level=2)

doc.add_paragraph(
    'The REST API connector (1715 lines) in ManifoldCF supports crawling any REST API endpoint. '
    'It includes 8 vendor presets and 5 pagination strategies.'
)

add_table(
    ['Vendor Preset', 'Default Endpoint', 'Auth Type', 'Pagination'],
    [
        ['Confluence', 'https://{domain}/wiki/rest/api/content', 'Bearer Token', 'OFFSET_LIMIT'],
        ['Jira', 'https://{domain}/rest/api/2/search', 'Basic Auth', 'OFFSET_LIMIT'],
        ['WordPress', 'https://{domain}/wp-json/wp/v2/posts', 'None/Bearer', 'PAGE_NUMBER'],
        ['GitHub', 'https://api.github.com/repos/{owner}/{repo}/contents', 'Bearer Token', 'LINK_HEADER'],
        ['SharePoint Online', 'https://{domain}/_api/web/lists/..', 'Bearer Token', 'CURSOR'],
        ['Notion', 'https://api.notion.com/v1/search', 'Bearer Token', 'CURSOR'],
        ['Drupal', 'https://{domain}/jsonapi/node/article', 'None/Basic', 'LINK_HEADER'],
        ['Alfresco REST', 'https://{domain}/alfresco/api/-default-/...', 'Basic Auth', 'OFFSET_LIMIT'],
    ]
)

doc.add_heading('6.6 ManifoldCF Init Scripts', level=2)

doc.add_paragraph(
    'The manifoldcf-init container runs a 6-step shell script (init-opensearch-output.sh) that '
    'auto-configures ManifoldCF when the Docker stack starts:'
)

add_code("""
#!/bin/sh
# init-opensearch-output.sh — 6-step MCF auto-configuration

# Step 0: Create OpenSearch index template for ACL fields
curl -X PUT "$OPENSEARCH_URL/_index_template/manifoldcf-acl" \\
  -H 'Content-Type: application/json' \\
  -d @opensearch-index-template.json

# Step 1: Create OpenSearch output connection
curl -X PUT "$MCF_API_URL/json/outputconnections/OpenSearch" \\
  -H 'Content-Type: application/json' \\
  -d '{"outputconnection":{"_children_":[
    {"_type_":"name","_value_":"OpenSearch"},
    {"_type_":"class_name","_value_":"...ElasticSearchConnector"},
    {"_type_":"configuration","_children_":[
      {"_type_":"SERVERLOCATION","_value_":"'$OPENSEARCH_URL'"},
      {"_type_":"INDEXNAME","_value_":"manifoldcf"},
      {"_type_":"INDEXTYPE","_value_":"_doc"}
    ]}
  ]}}'

# Step 2: Create CMIS repository connection
# ... (uses CMIS_* environment variables)

# Step 3: Create authorities sync job (optional)
# Step 4: Create document crawl job
# Step 5: Start the crawl job
""")

doc.add_heading('6.7 End-to-End Crawl Flow', level=2)

add_code("""
User Action                Admin Backend              ManifoldCF                 OpenSearch
─────────                  ──────────────             ──────────                 ──────────
Click "Create Source"  →   POST /source-systems   →   PUT /repositoryconnections →  (no action)
                           Save to PostgreSQL          Create CMIS/REST connection

Click "Start Sync"    →   POST /{id}/sync        →   POST /jobs                →  (no action)
                           Create & start job          Create crawl job
                                                      PUT /start/{jobId}
                                                      
                                                  →   Crawler runs:              →  Documents
                                                      - Connect to CMIS/REST        indexed with
                                                      - List documents              content +
                                                      - Extract content             metadata +
                                                      - Extract ACLs                ACL tokens
                                                      - Push to OpenSearch          

View Status           →   GET /{id}/status       →   GET /jobstatuses/{jobId}   →  (no action)
                           Refresh from MCF            Return live progress

Chat with RAG         →   POST /chat/completions  →  (no MCF call)              →  KNN search
                           Retrieval Advisor                                        on manifoldcf
                           searches OpenSearch                                      index
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  7. CHAT SYSTEM
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Chat System — How a User Message is Processed', level=1)

doc.add_heading('7.1 Overview & Two Chat Paths', level=2)

doc.add_paragraph(
    'The platform has two distinct chat execution paths that share the same model layer '
    'but differ in orchestration:'
)

add_table(
    ['Aspect', 'Admin Platform (BasicAgentExecutor)', 'Agent Framework (ReactAgent)'],
    [
        ['Entry Point', 'ChatController / AppChatController', 'AgentController (Studio)'],
        ['Orchestrator', 'BasicAgentExecutor (922 lines)', 'ReactAgent → CompiledGraph'],
        ['Tool Calling', 'Spring AI ToolCallingManager', 'AgentToolNode (custom impl)'],
        ['RAG', 'KnowledgeBaseRetrievalAdvisor', 'Custom Advisor or state-based'],
        ['Memory', 'MessageChatMemoryAdvisor', 'OverAllState + Checkpoint'],
        ['State', 'Stateless (per-request)', 'Stateful (OverAllState + persistence)'],
        ['Graph', 'No graph', 'Full StateGraph with conditions/loops'],
    ]
)

add_note(
    'BasicAgentExecutor in the admin platform does NOT use ReactAgent from the agent framework. '
    'They are completely separate systems. The admin platform uses Spring AI\'s ChatClient directly, '
    'while the agent framework compiles agents into state graphs.'
)

doc.add_heading('7.2 Request Entry Points (Controllers)', level=2)

add_code("""
// ChatController.java — Public API
@PostMapping("/api/v1/apps/chat/completions")
public SseEmitter chatCompletions(@RequestBody AgentRequest request) {
    // Public endpoint — uses published app config
    return agentService.streamChat(request, false);  // draft=false
}

// AppChatController.java — Console API (for testing drafts)
@PostMapping("/console/v1/apps/chat/completions")
public SseEmitter appChatCompletions(@RequestBody AgentRequest request) {
    // Console endpoint — uses draft app config
    return agentService.streamChat(request, true);   // draft=true
}
""")

doc.add_heading('7.3 AgentServiceImpl — Orchestrator', level=2)

add_code("""
// AgentServiceImpl.java
@Service
public class AgentServiceImpl implements AgentService {
    
    @Override
    public SseEmitter streamChat(AgentRequest request, boolean draft) {
        // 1. Validate request
        validateRequest(request);
        
        // 2. Load application configuration
        AppEntity app = appMapper.selectByAppId(request.getAppId());
        
        // 3. Load agent config (draft or published version)
        String configJson = draft ? app.getDraftConfig() : app.getPublishedConfig();
        AgentConfig config = JsonUtils.fromJson(configJson, AgentConfig.class);
        
        // 4. Build AgentContext
        AgentContext context = AgentContext.builder()
            .appId(request.getAppId())
            .config(config)
            .request(request)
            .workspaceId(getWorkspaceId())
            .memoryEnabled(config.getMemory() != null && config.getMemory().getEnabled())
            .build();
        
        // 5. Delegate to BasicAgentExecutor
        SseEmitter emitter = new SseEmitter(300_000L);  // 5 min timeout
        
        Flux<AgentResponse> responseFlux = basicAgentExecutor
            .streamExecute(context, request);
        
        // 6. Subscribe and forward SSE events
        responseFlux.subscribe(
            response -> emitter.send(response, MediaType.APPLICATION_JSON),
            error -> emitter.completeWithError(error),
            emitter::complete
        );
        
        return emitter;
    }
}
""")

doc.add_heading('7.4 BasicAgentExecutor — The Chat Engine', level=2)

doc.add_paragraph(
    'BasicAgentExecutor (922 lines) is the heart of the admin platform\'s chat system. It builds '
    'a Spring AI ChatClient with the appropriate model, tools, memory, and RAG advisors, then '
    'executes the chat with streaming and recursive tool call handling.'
)

doc.add_heading('7.4.1 AgentConfig Structure', level=3)

add_code("""
// AgentConfig.java — Stored as JSON in app.draft_config / app.published_config
public class AgentConfig {
    private String modelProvider;       // Provider code (e.g., "abc12345")
    private String model;               // Model ID (e.g., "qwen2.5:7b")
    private String instructions;        // System prompt
    private MemoryConfig memory;        // {enabled, windowSize}
    private ParameterConfig parameter;  // {temperature, topP, maxTokens}
    private List<String> tools;         // Tool callback names
    private List<String> mcpServers;    // MCP server names
    private FileSearchOptions fileSearch; // RAG config {kbIds, topK, threshold, ...}
    private Map<String, String> promptVariables;  // Template variables
    private String prologue;            // Opening message
}
""")

doc.add_heading('7.5 Building the ChatClient', level=2)

add_code("""
// BasicAgentExecutor.java — buildChatClient()
private ChatClient.Builder buildChatClient(AgentContext context, 
        ToolCallingChatOptions chatOptions,
        ToolCallbackProvider toolCallbackProvider, boolean enableTools) {
    
    AgentConfig config = context.getConfig();
    
    // 1. Get ChatModel from ModelFactory (OpenAI-compatible for ALL providers)
    ChatModel chatModel = modelFactory.getChatModel(
        config.getModelProvider(), config.getModel());
    
    // 2. Create base ChatClient
    ChatClient chatClient = ChatClient.builder(chatModel)
        .defaultAdvisors(new SimpleLoggerAdvisor())
        .build();
    
    ChatClient.Builder builder = chatClient.mutate();
    
    // 3. Add Memory Advisor (if enabled)
    if (context.isMemoryEnabled()) {
        MessageChatMemoryAdvisor memoryAdvisor = MessageChatMemoryAdvisor
            .builder(chatMemory).build();
        builder.defaultAdvisors(memoryAdvisor);
        
        String conversationId = String.format("%s_%s", 
            context.getAppId(), request.getConversationId());
        builder.defaultAdvisors(a -> a.param(CONVERSATION_ID, conversationId));
    }
    
    // 4. Add RAG Advisor (if knowledge bases configured)
    FileSearchOptions fileSearch = config.getFileSearch();
    if (fileSearch != null && fileSearch.getEnableSearch() 
            && !CollectionUtils.isEmpty(fileSearch.getKbIds())) {
        DocumentRetriever retriever = documentRetrieverManager
            .getDocumentRetriever(fileSearch);
        
        Advisor ragAdvisor = KnowledgeBaseRetrievalAdvisor.builder()
            .documentRetriever(retriever)
            .agentContext(context)
            .commonConfig(commonConfig)
            .build();
        builder.defaultAdvisors(ragAdvisor);
    }
    
    // 5. Add Tool Callbacks
    if (enableTools) {
        ToolCallback[] callbacks = toolCallbackProvider.getToolCallbacks();
        if (!ArrayUtils.isEmpty(callbacks)) {
            chatOptions.setToolCallbacks(Arrays.stream(callbacks).toList());
        }
    }
    
    return builder;
}
""")

doc.add_heading('7.6 Tool Execution', level=2)

add_code("""
// BasicAgentExecutor.java — processToolCallsRecursively()
private Flux<AgentResponse> processToolCallsRecursively(
        ChatClient.Builder clientBuilder, ChatResponse response, 
        Prompt prompt, ToolCallingManager toolCallingManager,
        ToolCallbackProvider provider, RequestContext requestContext,
        ToolCallingChatOptions chatOptions, boolean enableTools) {
    
    if (!response.hasToolCalls()) {
        // No tool calls — return the response directly
        return convertResponse(response, provider);
    }
    
    // Execute tool calls
    ToolExecutionResult result = toolCallingManager.executeToolCalls(
        prompt, response);
    
    // Build new prompt with tool results in conversation history
    Prompt newPrompt = new Prompt(
        result.conversationHistory(), chatOptions);
    
    // Recursively call the model with tool results
    return clientBuilder.build()
        .prompt(newPrompt)
        .options(chatOptions)
        .stream()
        .chatResponse()
        .concatMap(nextResponse -> processToolCallsRecursively(
            clientBuilder, nextResponse, newPrompt,
            toolCallingManager, provider, requestContext,
            chatOptions, enableTools));
}
""")

doc.add_heading('7.7 RAG Retrieval During Chat', level=2)

doc.add_paragraph(
    'When the ChatClient processes a request, the KnowledgeBaseRetrievalAdvisor intercepts '
    'the request before it reaches the LLM. The advisor retrieves relevant documents from '
    'OpenSearch and injects them into the system prompt.'
)

add_code("""
// Complete flow: User message → RAG → LLM → Response

// 1. User sends: "What is the company policy on remote work?"
// 2. KnowledgeBaseRetrievalAdvisor.adviseRequest() is called:
//    - Extracts query: "What is the company policy on remote work?"
//    - Calls DocumentRetriever.retrieve(query)
//    - DocumentRetriever creates OpenSearch KNN query:

// OpenSearch KNN Query (generated by Spring AI VectorStore):
{
  "query": {
    "knn": {
      "embedding": {
        "vector": [0.123, -0.456, ...],  // embedded query
        "k": 5
      }
    }
  }
}

// 3. Returns top-5 similar document chunks
// 4. Injects into system prompt:
//    "Use the following context to answer: {documents}"
// 5. LLM generates response using the context
""")

doc.add_heading('7.8 Streaming Response', level=2)

add_code("""
// BasicAgentExecutor.java — streamExecute()
@Override
public Flux<AgentResponse> streamExecute(AgentContext context, AgentRequest request) {
    AgentConfig config = context.getConfig();
    
    // Build all components
    ToolCallingChatOptions chatOptions = buildChatOptions(config);
    ToolCallbackProvider toolProvider = buildToolCallbackProvider(config, request);
    boolean enableTools = shouldEnableTools(request);
    List<Message> messages = buildMessages(context);
    ChatClient.Builder clientBuilder = buildChatClient(
        context, chatOptions, toolProvider, enableTools);
    
    // Create prompt and stream
    Prompt prompt = new Prompt(messages, chatOptions);
    
    return clientBuilder.build()
        .prompt(prompt)
        .options(chatOptions)
        .stream()
        .chatResponse()
        .concatMap(response -> processToolCallsRecursively(
            clientBuilder, response, prompt,
            ToolCallingManager.builder().build(),
            toolProvider, RequestContextHolder.getRequestContext(),
            chatOptions, enableTools));
}
""")

doc.add_heading('7.9 Complete Request Flow Diagram', level=2)

add_code("""
┌──────────┐  POST /chat/completions  ┌────────────────┐
│  Browser │─────────────────────────►│ ChatController  │
│  (React) │  SSE stream              │                │
│          │◄─────────────────────────│                │
└──────────┘                          └───────┬────────┘
                                              │
                                      ┌───────▼────────┐
                                      │AgentServiceImpl │
                                      │ validate()     │
                                      │ loadAppConfig()│
                                      │ buildContext() │
                                      └───────┬────────┘
                                              │
                                      ┌───────▼────────────────┐
                                      │ BasicAgentExecutor     │
                                      │                        │
                                      │ ┌──────────────────┐   │
                                      │ │ buildChatModel() │   │
                                      │ │ ModelFactory      │   │
                                      │ │ → OpenAiApi       │   │
                                      │ │ → OpenAiChatModel │   │
                                      │ └──────────────────┘   │
                                      │                        │
                                      │ ┌──────────────────┐   │
                                      │ │ buildChatClient() │   │
                                      │ │ + MemoryAdvisor   │   │
                                      │ │ + RAG Advisor     │◄──┼── KnowledgeBaseRetrievalAdvisor
                                      │ │ + ToolCallbacks   │   │     → DocumentRetriever
                                      │ └──────────────────┘   │     → OpenSearch KNN
                                      │                        │
                                      │ ┌──────────────────┐   │
                                      │ │ stream()          │   │
                                      │ │ → ChatClient      │   │
                                      │ │ → ChatModel.call()│──►│──► Ollama/DashScope/OpenAI
                                      │ │ → processTools() │   │
                                      │ └──────────────────┘   │
                                      └────────────────────────┘
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  8. AGENT FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Agent Framework (spring-ai-alibaba-agent-framework)', level=1)

doc.add_heading('8.1 Agent Hierarchy', level=2)

add_code("""
                        Agent (interface)
                           │
                      BaseAgent (abstract)
                      ├── name, description
                      ├── initGraph()        ← abstract
                      ├── compile()          ← calls initGraph() → StateGraph → compile()
                      ├── stream()           ← delegates to CompiledGraph
                      ├── invoke()
                      │
               ┌──────┴──────────┐
               │                 │
          ReactAgent         FlowAgent (abstract)
          │ llmNode           │
          │ toolNode          ├── SequentialAgent
          │ hooks             ├── ParallelAgent
          │ interceptors      ├── LlmRoutingAgent
          │                   └── LoopAgent
          │
          └── initGraph() creates:
              START → [hooks] → AgentLlmNode ←→ AgentToolNode → [hooks] → END
""")

doc.add_heading('8.2 ReactAgent — The Core Agent', level=2)

doc.add_paragraph(
    'ReactAgent implements the ReAct (Reasoning + Acting) pattern. It creates a StateGraph '
    'with two primary nodes: AgentLlmNode (model invocation) and AgentToolNode (tool execution), '
    'connected by conditional edges that route based on whether the model requested tool calls.'
)

add_code("""
// ReactAgent.java — Building an agent
ReactAgent agent = ReactAgent.builder()
    .name("customer_support")
    .model(chatModel)                        // Spring AI ChatModel
    .instruction("You are a helpful customer support agent.")
    .tools(searchTool, emailTool)            // ToolCallback instances
    .hooks(
        new SummarizationHook(chatModel),    // Compress long conversations
        new HumanInTheLoopHook()             // Require approval for sensitive ops
    )
    .saver(new MemorySaver())                // Checkpoint persistence
    .build();

// Execute the agent
agent.stream(Map.of("input", "How do I reset my password?"), runnableConfig)
    .subscribe(output -> System.out.println(output));
""")

doc.add_heading('8.2.1 Graph Construction (initGraph)', level=3)

add_code("""
// ReactAgent.java — initGraph()
@Override
protected StateGraph initGraph() throws GraphStateException {
    // Create state graph with custom key strategies for messages
    StateGraph graph = new StateGraph(name, 
        buildMessagesKeyStrategyFactory(hooks), stateSerializer);
    
    // Add the two core nodes
    graph.addNode(AGENT_MODEL_NAME, node_async(this.llmNode));
    if (hasTools) {
        graph.addNode(AGENT_TOOL_NAME, node_async(this.toolNode));
    }
    
    // Wire hook nodes (before/after agent, before/after model)
    List<Hook> beforeAgentHooks = filterHooksByPosition(hooks, BEFORE_AGENT);
    List<Hook> afterAgentHooks = filterHooksByPosition(hooks, AFTER_AGENT);
    List<Hook> beforeModelHooks = filterHooksByPosition(hooks, BEFORE_MODEL);
    List<Hook> afterModelHooks = filterHooksByPosition(hooks, AFTER_MODEL);
    
    // Connect: START → entry → model → tools → model → exit → END
    String entryNode = determineEntryNode(beforeAgentHooks, beforeModelHooks);
    graph.addEdge(START, entryNode);
    
    // Conditional edge: Model → Tools (if tool calls) or End (if no tool calls)
    setupToolRouting(graph, loopExitNode, loopEntryNode, exitNode, this);
    
    return graph;
}
""")

doc.add_heading('8.2.2 Tool Routing Logic', level=3)

add_code("""
// ReactAgent.java — makeModelToTools()
// This EdgeAction decides: route to tools or to end?
private EdgeAction makeModelToTools(String modelDest, String endDest) {
    return state -> {
        List<Message> messages = (List<Message>) state.value("messages").orElse(List.of());
        Message lastMessage = messages.get(messages.size() - 1);
        
        if (lastMessage instanceof AssistantMessage assistant) {
            if (assistant.hasToolCalls()) {
                return AGENT_TOOL_NAME;  // → AgentToolNode
            } else {
                return endDest;          // → END (or after-model hooks)
            }
        }
        return endDest;
    };
}

// ReactAgent.java — makeToolsToModelEdge()
// After tools execute, route back to model for next reasoning step
private EdgeAction makeToolsToModelEdge(String modelDest, String endDest) {
    return state -> {
        // Check for return_direct tools (skip model, go directly to output)
        // Otherwise, route back to model node for next iteration
        return modelDest;  // → AgentLlmNode (ReAct loop)
    };
}
""")

doc.add_heading('8.3 AgentLlmNode — Model Invocation', level=3)

add_code("""
// AgentLlmNode.java — apply() (simplified)
@Override
public Map<String, Object> apply(OverAllState state, RunnableConfig config) {
    // 1. Extract messages from state
    List<Message> messages = (List<Message>) state.value("messages").get();
    
    // 2. Build ModelRequest with tools, system prompt, options
    ModelRequest request = ModelRequest.builder()
        .messages(messages)
        .options(chatOptions)
        .build();
    
    // 3. Apply interceptor chain (before/after model call)
    ModelCallHandler baseHandler = req -> {
        Flux<ChatResponse> flux = buildChatClientRequestSpec(req, config)
            .stream().chatResponse();
        return ModelResponse.of(flux);
    };
    ModelCallHandler chained = InterceptorChain
        .chainModelInterceptors(modelInterceptors, baseHandler);
    
    // 4. Execute and return updated state
    ModelResponse response = chained.call(request);
    return Map.of("messages", response.getMessage());
}
""")

doc.add_heading('8.4 AgentToolNode — Tool Execution', level=3)

add_code("""
// AgentToolNode.java — apply() (simplified)
@Override
public Map<String, Object> apply(OverAllState state, RunnableConfig config) {
    List<Message> messages = (List<Message>) state.value("messages").get();
    AssistantMessage lastMsg = (AssistantMessage) messages.get(messages.size() - 1);
    
    List<AssistantMessage.ToolCall> toolCalls = lastMsg.getToolCalls();
    
    if (parallelToolExecution && toolCalls.size() > 1) {
        return executeToolCallsParallel(toolCalls, state, config);
    } else {
        return executeToolCallsSequential(toolCalls, state, config);
    }
}

private Map<String, Object> executeToolCallsSequential(
        List<AssistantMessage.ToolCall> toolCalls, ...) {
    List<ToolResponse> responses = new ArrayList<>();
    
    for (AssistantMessage.ToolCall toolCall : toolCalls) {
        // Find matching ToolCallback by name
        ToolCallback callback = findCallback(toolCall.name());
        
        // Apply interceptor chain (before/after tool execution)
        ToolCallHandler baseHandler = req -> {
            String result = callback.call(req.getArguments(), 
                new ToolContext(toolContextMap));
            return ToolCallResponse.of(req.getToolCallId(), 
                req.getToolName(), result);
        };
        ToolCallHandler chained = InterceptorChain
            .chainToolInterceptors(toolInterceptors, baseHandler);
        
        ToolCallResponse response = chained.call(toolCallRequest);
        responses.add(response.toToolResponse());
    }
    
    ToolResponseMessage msg = ToolResponseMessage.builder()
        .responses(responses).build();
    return Map.of("messages", msg);
}
""")

doc.add_heading('8.5 Multi-Agent Patterns', level=2)

add_table(
    ['Pattern', 'Class', 'Description', 'Use Case'],
    [
        ['Sequential', 'SequentialAgent', 'Agents run in order, output of one feeds into next', 'Pipeline processing (research → write → review)'],
        ['Parallel', 'ParallelAgent', 'Agents run concurrently, results merged via MergeStrategy', 'Independent tasks (search web + search DB simultaneously)'],
        ['LLM Routing', 'LlmRoutingAgent', 'LLM decides which sub-agent handles the request', 'Intent classification (route to billing vs support vs sales)'],
        ['Loop', 'LoopAgent', 'Agent runs repeatedly until condition met', 'Iterative refinement (edit until quality threshold met)'],
    ]
)

add_code("""
// SequentialAgent — Pipeline execution
SequentialAgent pipeline = SequentialAgent.builder()
    .name("research_pipeline")
    .subAgents(researchAgent, writerAgent, reviewerAgent)
    .build();

// ParallelAgent — Concurrent execution
ParallelAgent parallel = ParallelAgent.builder()
    .name("multi_search")
    .subAgents(webSearchAgent, dbSearchAgent, cacheSearchAgent)
    .mergeStrategy(MergeStrategy.CONCATENATE)
    .build();

// LlmRoutingAgent — LLM-based routing
LlmRoutingAgent router = LlmRoutingAgent.builder()
    .name("support_router")
    .model(chatModel)
    .subAgents(billingAgent, technicalAgent, salesAgent)
    .build();

// LoopAgent — Iterative execution
LoopAgent loop = LoopAgent.builder()
    .name("refiner")
    .subAgent(editorAgent)
    .loopStrategy(LoopStrategy.CONDITION)
    .condition(state -> state.value("quality_score")
        .map(s -> (Double) s < 0.9).orElse(true))
    .maxIterations(5)
    .build();
""")

doc.add_heading('8.6 Hook System (Context Engineering)', level=2)

doc.add_paragraph(
    'Hooks allow injecting logic at specific points in the agent execution cycle. They are the '
    'primary mechanism for context engineering — modifying the agent\'s behavior without changing '
    'its core logic.'
)

add_table(
    ['Hook', 'Position', 'Purpose'],
    [
        ['SummarizationHook', 'AFTER_MODEL', 'Compresses long message histories to stay within token limits'],
        ['HumanInTheLoopHook', 'BEFORE_AGENT', 'Pauses execution for human approval before proceeding'],
        ['ShellToolAgentHook', 'BEFORE_AGENT + AFTER_AGENT', 'Manages shell session lifecycle for ShellTool'],
    ]
)

add_code("""
// Hook positions in the graph:
// START → [BEFORE_AGENT hooks] → [BEFORE_MODEL hooks] → AgentLlmNode
//     → AgentToolNode → [AFTER_MODEL hooks] → AgentLlmNode (loop)
//     → [AFTER_AGENT hooks] → END
""")

doc.add_heading('8.7 Interceptor Chain', level=2)

add_code("""
// InterceptorChain.java — Wraps model/tool calls in interceptor layers
public static ModelCallHandler chainModelInterceptors(
        List<ModelInterceptor> interceptors, ModelCallHandler baseHandler) {
    
    ModelCallHandler current = baseHandler;
    // Wrap from last to first (outermost interceptor runs first)
    for (int i = interceptors.size() - 1; i >= 0; i--) {
        ModelInterceptor interceptor = interceptors.get(i);
        ModelCallHandler next = current;
        current = request -> {
            // Before
            ModelRequest modified = interceptor.beforeModelCall(request);
            // Execute
            ModelResponse response = next.call(modified);
            // After
            return interceptor.afterModelCall(modified, response);
        };
    }
    return current;
}
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  9. GRAPH CORE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('9. Graph Core (spring-ai-alibaba-graph-core)', level=1)

doc.add_paragraph(
    'The Graph Core module is the low-level workflow engine that powers the agent framework. '
    'It provides StateGraph (workflow definition), CompiledGraph (execution engine), OverAllState '
    '(state container), and a checkpoint/persistence system.'
)

doc.add_heading('9.1 StateGraph — Workflow Definition', level=2)

add_code("""
// StateGraph.java — Core workflow definition API
public class StateGraph {
    public static final String START = "__START__";
    public static final String END   = "__END__";
    public static final String ERROR = "__ERROR__";
    
    final Nodes nodes = new Nodes();              // LinkedHashSet<Node>
    final Edges edges = new Edges();              // LinkedList<Edge>
    
    // Node operations
    public void addNode(String id, AsyncNodeActionWithConfig action)
    public void addNode(String id, CompiledGraph subGraph)  // compiled subgraph
    public void addNode(String id, StateGraph subGraph)      // uncompiled subgraph
    
    // Edge operations
    public void addEdge(String source, String target)         // direct edge
    public void addEdge(String source, List<String> targets)  // parallel (fan-out)
    public void addEdge(List<String> sources, String target)  // fan-in
    
    // Conditional edges
    public void addConditionalEdges(String source, AsyncCommandAction condition,
                                    Map<String, String> mappings)
    
    // Multi-routing conditional edges
    public void addConditionalEdges(String source, AsyncMultiCommandAction condition,
                                    Map<String, String> mappings)
    
    // Compile to executable graph
    public CompiledGraph compile(CompileConfig config)
    public CompiledGraph compile()  // uses MemorySaver default
}
""")

doc.add_heading('9.1.1 Example: Building a Custom Workflow', level=3)

add_code("""
// Building a custom classification + processing workflow
StateGraph graph = new StateGraph("document_processor", keyStrategyFactory);

// Add nodes
graph.addNode("classifier", node_async(state -> {
    String input = state.value("input").map(Object::toString).orElse("");
    String type = classifyDocument(input);  // "invoice", "receipt", "contract"
    return Map.of("doc_type", type);
}));

graph.addNode("invoice_processor", node_async(state -> {
    // Process invoice...
    return Map.of("result", processInvoice(state));
}));

graph.addNode("receipt_processor", node_async(state -> {
    return Map.of("result", processReceipt(state));
}));

graph.addNode("contract_processor", node_async(state -> {
    return Map.of("result", processContract(state));
}));

// Wire edges
graph.addEdge(START, "classifier");
graph.addConditionalEdges("classifier", 
    edge_async(state -> state.value("doc_type").map(Object::toString).orElse("receipt")),
    Map.of(
        "invoice",  "invoice_processor",
        "receipt",  "receipt_processor",
        "contract", "contract_processor"
    )
);
graph.addEdge("invoice_processor", END);
graph.addEdge("receipt_processor", END);
graph.addEdge("contract_processor", END);

// Compile and run
CompiledGraph compiled = graph.compile();
compiled.invoke(Map.of("input", "Invoice #12345 for $500"));
""")

doc.add_heading('9.2 CompiledGraph — Execution Engine', level=2)

add_code("""
// CompiledGraph.java — Key execution methods
public class CompiledGraph {
    public final StateGraph stateGraph;
    public final CompileConfig compileConfig;
    final Map<String, Node.ActionFactory> nodeFactories;  // thread-safe
    final Map<String, EdgeValue> edges;
    private int maxIterations = 25;
    
    // Reactive streaming
    public Flux<NodeOutput> stream(Map<String, Object> inputs, RunnableConfig config)
    public Flux<NodeOutput> streamSnapshots(Map<String, Object> inputs, RunnableConfig config)
    
    // Blocking invocation
    public Optional<OverAllState> invoke(Map<String, Object> inputs, RunnableConfig config)
    
    // State management
    public RunnableConfig updateState(RunnableConfig config, 
        Map<String, Object> values, String asNode)
    public Optional<StateSnapshot> stateOf(RunnableConfig config)
    public Collection<StateSnapshot> getStateHistory(RunnableConfig config)
}
""")

doc.add_paragraph(
    'The compilation pipeline transforms a StateGraph into an executable CompiledGraph:'
)

add_code("""
// Compilation Pipeline:
// 1. StateGraph.compile(config) calls:
//    → validateGraph()       - ensures START edge exists, all nodes valid
//    → new CompiledGraph(this, config)
//
// 2. CompiledGraph constructor:
//    → processNodesAndEdges()  - flatten subgraphs, resolve edges
//    → Store node ActionFactories (not instances!) for thread safety
//    → Process edges into three categories:
//       a) Direct edge: source → target
//       b) Conditional edge: source → condition() → mapping → target
//       c) Parallel edge: source → ParallelNode(targets) → common_target
//
// 3. Execution (stream/invoke):
//    → Get entry point from START edge
//    → Loop: execute node → determine next → check interrupts → checkpoint
//    → Max iterations: 25 (configurable)
//    → Each node execution:
//       a) ActionFactory creates fresh action instance (thread-safe)
//       b) Action.apply(state, config) returns Map<String, Object>
//       c) State updated via KeyStrategy merge
//       d) Checkpoint saved via BaseCheckpointSaver
""")

doc.add_heading('9.3 OverAllState — State Container', level=2)

add_code("""
// OverAllState.java — Core state management
public final class OverAllState implements Serializable {
    public static final Object MARK_FOR_REMOVAL = new Object();
    public static final String DEFAULT_INPUT_KEY = "input";
    
    private final Map<String, Object> data;
    private final Map<String, KeyStrategy> keyStrategies;
    
    // Input merge — uses registered strategies per key
    public OverAllState mergeWith(Map<String, Object> partialState, 
                                  Map<String, KeyStrategy> strategies) {
        for (Map.Entry<String, Object> entry : partialState.entrySet()) {
            String key = entry.getKey();
            Object newValue = entry.getValue();
            
            if (newValue == MARK_FOR_REMOVAL) {
                data.remove(key);  // Special sentinel: delete key
                continue;
            }
            
            KeyStrategy strategy = strategies.getOrDefault(key, KeyStrategy.REPLACE);
            Object oldValue = data.get(key);
            data.put(key, strategy.apply(oldValue, newValue));
        }
        return this;
    }
    
    // Typed getters
    public Optional<Object> value(String key)
    public <T> Optional<T> value(String key, TypeRef<T> typeRef)
    public OverAllState cloneState()   // deep copy
}
""")

doc.add_heading('9.4 Node & Edge Model', level=2)

add_code("""
// Node.java — Graph node with lazy action factory
public class Node {
    private final String id;
    private final ActionFactory actionFactory;
    
    // ActionFactory creates fresh action per execution (thread-safe)
    public interface ActionFactory {
        AsyncNodeActionWithConfig apply(CompileConfig config);
    }
}

// Edge.java — Graph edge with optional parallelism
public record Edge(String sourceId, List<EdgeValue> targets) {
    public boolean isParallel() { return targets.size() > 1; }
}

// EdgeValue.java — Edge target with optional condition
public record EdgeValue(String id, EdgeCondition value) {
    public EdgeValue(String id)          // direct: always go to id
    public EdgeValue(EdgeCondition value) // conditional: evaluate at runtime
}

// EdgeCondition.java — Routing logic
public record EdgeCondition(Object action, Map<String, String> mappings) {
    // Single routing: condition returns one key → lookup in mappings
    public static EdgeCondition single(AsyncCommandAction action, 
                                       Map<String, String> mappings)
    // Multi routing: condition returns multiple keys → parallel execution
    public static EdgeCondition multi(AsyncMultiCommandAction action, 
                                      Map<String, String> mappings)
}
""")

doc.add_heading('9.5 KeyStrategy — State Merge Policies', level=2)

add_table(
    ['Strategy', 'Behavior', 'Use Case'],
    [
        ['ReplaceStrategy', 'New value replaces old value', 'Simple values (strings, numbers, current status)'],
        ['AppendStrategy', 'New value appended to list; supports dedup and removal', 'Message history, accumulated results'],
        ['MergeStrategy', 'Deep-merge Maps; throws for incompatible types', 'Configuration objects, metadata'],
    ]
)

add_code("""
// KeyStrategy interface
public interface KeyStrategy extends BiFunction<Object, Object, Object> {
    KeyStrategy REPLACE = new ReplaceStrategy();  // (old, new) → new
    KeyStrategy APPEND  = new AppendStrategy();   // (old, new) → [...old, new]
    KeyStrategy MERGE   = new MergeStrategy();    // (old, new) → {..old, ...new}
}

// Example: Messages use APPEND so conversation history accumulates
// "messages" key → AppendStrategy
// Each node returns Map.of("messages", newMessage)
// State automatically appends to existing message list
""")

doc.add_heading('9.6 Checkpoint & Persistence', level=2)

add_code("""
// BaseCheckpointSaver.java — Persistence interface
public interface BaseCheckpointSaver {
    Collection<Checkpoint> list(RunnableConfig config);
    Optional<Checkpoint> get(RunnableConfig config);
    RunnableConfig put(RunnableConfig config, Checkpoint checkpoint);
}

// Checkpoint.java — State snapshot
public class Checkpoint {
    private final String id;           // UUID
    private Map<String, Object> state; // full state snapshot
    private String nodeId;             // which node created this
    private String nextNodeId;         // what node to run next
}

// Available implementations:
// MemorySaver      — In-memory (development, testing)
// PostgresSaver    — PostgreSQL (production)
// MysqlSaver       — MySQL (production)
// RedisSaver       — Redis (production)
// MongoSaver       — MongoDB (production)
// OracleSaver      — Oracle (production)
// FileSystemSaver  — File system (edge deployment)
""")

doc.add_heading('9.7 Compilation Pipeline', level=2)

add_code("""
// Full compilation flow:

// 1. User code:
StateGraph graph = new StateGraph("myGraph", keyFactory);
graph.addNode("nodeA", action);
graph.addEdge(START, "nodeA");
graph.addEdge("nodeA", END);
CompiledGraph compiled = graph.compile();

// 2. Internal: StateGraph.compile()
//    → validateGraph()
//       - Every Node validates: not blank, not START/END, no "__" prefix
//       - Ensures START edge exists
//       - Every Edge validates: source/target nodes exist
//    → new CompiledGraph(this, config)

// 3. Internal: CompiledGraph constructor
//    → processNodesAndEdges(stateGraph)
//       - Flattens subgraphs (both compiled and uncompiled)
//       - Renames subgraph internal nodes to avoid collisions
//       - Merges interrupt configs from subgraphs
//    → For each node: store ActionFactory (lazy, thread-safe)
//    → For each edge:
//       - Direct → EdgeValue(targetId)
//       - Conditional → EdgeValue(EdgeCondition)
//       - Parallel → ParallelNode wrapping all targets
//    → Store maxIterations (default 25)

// 4. Execution: compiled.stream(inputs)
//    → Create initial OverAllState from inputs
//    → Get START edge target → first node
//    → Loop:
//       a. Check interrupt-before for current node
//       b. ActionFactory.apply(config) → create action instance
//       c. action.apply(state, runnableConfig) → partial state updates
//       d. Merge updates into OverAllState via KeyStrategies
//       e. Create Checkpoint (nodeId, state snapshot, nextNodeId)
//       f. Save checkpoint via BaseCheckpointSaver
//       g. Check interrupt-after for current node
//       h. Determine next node via edge routing
//       i. If next == END or iterations >= maxIterations → stop
//       j. Emit NodeOutput to Flux subscriber
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  10. OBSERVABILITY & TRACING
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('10. Observability & Tracing', level=1)

doc.add_heading('10.1 Micrometer Observation API', level=2)

doc.add_paragraph(
    'The platform uses Micrometer Observation API (not raw OpenTelemetry) which bridges to '
    'both Micrometer metrics AND distributed tracing backends (OTLP, Zipkin, Jaeger). The '
    'graph observation module provides three-level tracing for graph execution.'
)

doc.add_heading('10.2 Graph-Level Observations', level=2)

add_table(
    ['Observation Kind', 'Metric Name', 'Attributes'],
    [
        ['GRAPH', 'spring.ai.alibaba.graph', 'graph_name, is_error, gen_ai.input, gen_ai.output'],
        ['NODE', 'spring.ai.alibaba.graph.node', 'graph_name, node_name, is_error, langfuse.input/output'],
        ['EDGE', 'spring.ai.alibaba.graph.edge', 'graph_name, edge_name, is_error'],
    ]
)

add_code("""
// GraphObservationLifecycleListener.java
// Creates parent-child observation spans:

Graph Observation (parent)                   ← whole graph execution
  ├── Node "AgentLlmNode" Observation (child)  ← model invocation
  │     └── scope opened for context propagation
  ├── Node "AgentToolNode" Observation (child) ← tool execution
  │     └── scope opened for context propagation
  └── ...

// Each node observation:
// 1. Opens a scope for distributed trace context propagation
// 2. Records state as high-cardinality attributes (truncated to 1000 chars)
// 3. Records gen_ai.input / gen_ai.output for LLM observability
""")

doc.add_heading('10.3 Metrics & Counters', level=2)

add_code("""
// Auto-configured observation handlers create Micrometer Counters:
// - spring.ai.alibaba.graph        → graph execution count
// - spring.ai.alibaba.graph.node   → node execution count  
// - spring.ai.alibaba.graph.edge   → edge transition count

// SpringAiAlibabaChatModelObservationConvention
// Extends Spring AI's ChatModelObservationConvention:
// - Adds gen_ai.content.input  (from user messages)
// - Adds gen_ai.content.output (from assistant messages)
// - Adds langfuse.input / langfuse.output (Langfuse-compatible)
""")

doc.add_heading('10.4 Auto-Configuration', level=2)

add_code("""
// GraphObservationAutoConfiguration.java
// Registered via: META-INF/spring/org.springframework.boot.autoconfigure.AutoConfiguration.imports
// Properties: spring.ai.alibaba.graph.observation.enabled (default: true)

@AutoConfiguration
@ConditionalOnProperty(prefix = "spring.ai.alibaba.graph.observation", 
                       name = "enabled", havingValue = "true", matchIfMissing = true)
public class GraphObservationAutoConfiguration {
    
    @Bean  // Micrometer context-propagation (if on classpath)
    ObservationRegistryCustomizer<?> propagationCustomizer() { ... }
    
    @Bean  // Graph execution tracing
    GraphObservationLifecycleListener graphLifecycleListener() { ... }
    
    @Bean  // Spring AI chat model convention (GenAI attributes)
    SpringAiAlibabaChatModelObservationConvention chatModelConvention() { ... }
    
    @Bean  // Default ObservationHandlerProvider
    ObservationHandlerProvider observationHandlerProvider() { ... }
    
    @Bean  // Counter: spring.ai.alibaba.graph
    GraphObservationHandler graphObservationHandler(MeterRegistry registry) { ... }
    
    @Bean  // Counter: spring.ai.alibaba.graph.node
    GraphNodeObservationHandler nodeObservationHandler(MeterRegistry registry) { ... }
    
    @Bean  // Counter: spring.ai.alibaba.graph.edge
    GraphEdgeObservationHandler edgeObservationHandler(MeterRegistry registry) { ... }
}
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  11. STUDIO MODULE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('11. Studio Module (Debug UI)', level=1)

doc.add_paragraph(
    'The Studio module (spring-ai-alibaba-studio) provides an embedded debugging UI for '
    'visualizing and running agents. It ships as a Spring Boot auto-configuration module with '
    'a React frontend (agent-chat-ui).'
)

add_code("""
// Studio Architecture:

┌─────────────────────────────┐
│  React Chat UI              │ (localhost:3000 or 3001)
│  agent-chat-ui/             │
│  ├── Agent list sidebar     │
│  ├── Chat interface         │
│  └── Thread management      │
└──────────┬──────────────────┘
           │ HTTP / SSE
┌──────────▼──────────────────┐
│  Spring Boot Auto-Config    │
│  StudioAutoConfiguration    │
│                             │
│  ┌────────────────────┐     │
│  │ AgentController    │     │  GET /list-apps → lists registered agents
│  │ /api/v1/agents     │     │
│  ├────────────────────┤     │
│  │ ExecutionController│     │  POST /run_sse → SSE streaming execution
│  │                    │     │  POST /resume_sse → resume after interrupt
│  ├────────────────────┤     │
│  │ ThreadController   │     │  CRUD for conversation threads
│  ├────────────────────┤     │
│  │ AgentLoader (user) │     │  User implements: load(name) + listNames()
│  └────────────────────┘     │
└─────────────────────────────┘
""")

add_code("""
// User implements AgentLoader to register their agents:
@Component
public class MyAgentLoader implements AgentLoader {
    
    @Override
    public List<String> listNames() {
        return List.of("customer_support", "code_reviewer", "researcher");
    }
    
    @Override
    public BaseAgent load(String name) {
        return switch (name) {
            case "customer_support" -> customerSupportAgent;
            case "code_reviewer" -> codeReviewerAgent;
            case "researcher" -> researcherAgent;
            default -> throw new IllegalArgumentException("Unknown agent: " + name);
        };
    }
}

// SSE Execution Flow:
// 1. Client sends POST /run_sse with {agentName, input, threadId}
// 2. ExecutionController loads agent via AgentLoader
// 3. Builds RunnableConfig with threadId and metadata
// 4. Calls agent.stream(inputs, config) → Flux<NodeOutput>
// 5. Each NodeOutput emitted as SSE event to client
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  12. FRONTEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('12. Frontend Architecture', level=1)

doc.add_heading('12.1 Technology Stack', level=2)

add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['React', '18', 'UI framework'],
        ['UmiJS', '4', 'Application framework (routing, build, plugins)'],
        ['TypeScript', '5', 'Type safety'],
        ['Ant Design', '5', 'UI component library'],
        ['ahooks', '3', 'React hooks library'],
        ['XFlow', '-', 'Visual workflow editor (spark-flow package)'],
        ['Zustand', '-', 'State management (spark-flow)'],
        ['Node.js', '>= 20', 'Runtime'],
    ]
)

doc.add_heading('12.2 Package Structure', level=2)

add_code("""
frontend/
├── packages/
│   ├── main/              ← Main workbench application
│   │   ├── src/
│   │   │   ├── pages/     ← Page components
│   │   │   │   ├── Agent/         ← Agent creation & management
│   │   │   │   ├── Chat/          ← Chat interface
│   │   │   │   ├── Knowledge/     ← Knowledge base management
│   │   │   │   ├── Source/        ← Source system management
│   │   │   │   ├── Settings/      ← Model service configuration
│   │   │   │   ├── MCP/           ← MCP server management
│   │   │   │   └── Plugin/        ← Plugin management
│   │   │   ├── services/  ← API service layer
│   │   │   ├── layouts/   ← Layout components (SideMenuLayout)
│   │   │   └── request.ts ← Axios HTTP client
│   │   ├── .umirc.ts      ← Routes, proxy config
│   │   └── .env            ← Environment variables
│   │
│   ├── spark-flow/        ← Visual workflow editor
│   │   ├── src/
│   │   │   ├── core/      ← XFlow graph engine
│   │   │   ├── nodes/     ← Custom node types
│   │   │   └── store/     ← Zustand state management
│   │
│   └── spark-i18n/        ← Internationalization
│       └── src/
│           ├── zh-CN/     ← Chinese translations
│           └── en-US/     ← English translations
""")

doc.add_heading('12.3 Routing & Navigation', level=2)

add_code("""
// .umirc.ts — Route configuration
export default {
  routes: [
    { path: '/', redirect: '/agent' },
    { path: '/agent', component: '@/pages/Agent' },
    { path: '/agent/create', component: '@/pages/Agent/Create' },
    { path: '/agent/:id', component: '@/pages/Agent/Detail' },
    { path: '/chat', component: '@/pages/Chat' },
    { path: '/knowledge', component: '@/pages/Knowledge' },
    { path: '/knowledge/create', component: '@/pages/Knowledge/Create' },
    { path: '/source', component: '@/pages/Source' },
    { path: '/source/create', component: '@/pages/Source/Create' },
    { path: '/mcp', component: '@/pages/MCP' },
    { path: '/plugin', component: '@/pages/Plugin' },
    { path: '/settings/model-service', component: '@/pages/Settings/ModelService' },
  ],
  proxy: {
    '/api': { target: 'http://127.0.0.1:8080', changeOrigin: true },
    '/console': { target: 'http://127.0.0.1:8080', changeOrigin: true },
  }
}
""")

doc.add_heading('12.4 Key Pages', level=2)

add_table(
    ['Page', 'Path', 'Description'],
    [
        ['Agent List', '/agent', 'List/create/manage AI agents with tools, knowledge, and instructions'],
        ['Agent Builder', '/agent/create', 'Visual agent configuration with model selection, tools, and testing'],
        ['Chat', '/chat', 'Interactive chat interface with streaming, tool calls, and file references'],
        ['Knowledge Base', '/knowledge', 'Create/manage knowledge bases, upload documents, test retrieval'],
        ['Source System', '/source', 'Connect external repositories (CMIS, REST API), manage crawl jobs'],
        ['Source Create', '/source/create', 'Two-step vendor preset selection + connection configuration'],
        ['Model Service', '/settings/model-service', 'Add/configure model providers and individual models'],
        ['MCP Servers', '/mcp', 'Manage Model Context Protocol servers'],
        ['Plugins', '/plugin', 'Manage tool plugins'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  13. APPENDIX — FILE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('13. Appendix — File Reference', level=1)

doc.add_paragraph('Key files referenced in this document:')

files = [
    ['Docker Infrastructure', '', ''],
    ['', 'spring-ai-alibaba-admin/docker/middleware/docker-compose-arm.yaml', 'Docker Compose orchestration (13+ services)'],
    ['', 'spring-ai-alibaba-admin/docker/middleware/manifoldcf/config/init-opensearch-output.sh', '6-step MCF auto-configuration script'],
    ['', 'spring-ai-alibaba-admin/docker/middleware/ollama/models.conf', 'Models to auto-pull into Ollama'],
    ['', '', ''],
    ['Admin Backend — Controllers', '', ''],
    ['', 'admin-server-start/.../controller/ChatController.java', 'Public chat API'],
    ['', 'admin-server-start/.../controller/AppChatController.java', 'Console chat API (draft testing)'],
    ['', 'admin-server-start/.../controller/KnowledgeBaseController.java', 'Knowledge base CRUD (7 endpoints)'],
    ['', 'admin-server-start/.../controller/DocumentController.java', 'Document CRUD (8 endpoints)'],
    ['', 'admin-server-start/.../controller/SourceSystemController.java', 'Source system CRUD (11 endpoints)'],
    ['', 'admin-server-start/.../controller/ProviderController.java', 'Provider + model CRUD'],
    ['', 'admin-server-start/.../controller/ModelController.java', 'Model selector API'],
    ['', '', ''],
    ['Admin Backend — Services', '', ''],
    ['', 'admin-server-runtime/.../service/AgentServiceImpl.java', 'Chat orchestrator'],
    ['', 'admin-server-runtime/.../service/BasicAgentExecutor.java', 'Chat engine (922 lines)'],
    ['', 'admin-server-runtime/.../service/KnowledgeBaseServiceImpl.java', 'KB CRUD + OpenSearch index creation'],
    ['', 'admin-server-runtime/.../service/DocumentServiceImpl.java', 'Document CRUD + MQ publishing'],
    ['', 'admin-server-runtime/.../service/DocumentIndexHandler.java', 'MQ consumer — parse/chunk/embed/store'],
    ['', 'admin-server-runtime/.../service/KnowledgeBaseIndexPipeline.java', 'Document processing pipeline'],
    ['', 'admin-server-runtime/.../service/SourceSystemServiceImpl.java', 'Source system CRUD + MCF bridge'],
    ['', 'admin-server-runtime/.../service/ManifoldCFBridgeService.java', 'MCF REST API client'],
    ['', 'admin-server-runtime/.../service/ModelFactory.java', 'Universal model factory (OpenAI-compatible)'],
    ['', 'admin-server-runtime/.../service/ProviderManager.java', 'Provider CRUD + Redis cache'],
    ['', 'admin-server-runtime/.../service/ModelManager.java', 'Model CRUD + auto-sync'],
    ['', '', ''],
    ['Admin Backend — RAG', '', ''],
    ['', 'admin-server-runtime/.../retriever/DocumentRetrieverManager.java', 'Parallel multi-KB retriever factory'],
    ['', 'admin-server-runtime/.../retriever/KnowledgeBaseDocumentRetriever.java', 'Per-KB vector search'],
    ['', 'admin-server-runtime/.../advisor/KnowledgeBaseRetrievalAdvisor.java', 'Spring AI Advisor for RAG'],
    ['', 'admin-server-runtime/.../service/OpenSearchVectorStoreService.java', 'OpenSearch KNN index management'],
    ['', '', ''],
    ['Agent Framework', '', ''],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/ReactAgent.java', 'ReAct agent (1163 lines)'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/AgentLlmNode.java', 'LLM invocation node (622 lines)'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/AgentToolNode.java', 'Tool execution node (997 lines)'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/SequentialAgent.java', 'Sequential multi-agent'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/ParallelAgent.java', 'Parallel multi-agent'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/LlmRoutingAgent.java', 'LLM-based routing agent'],
    ['', 'spring-ai-alibaba-agent-framework/.../agent/LoopAgent.java', 'Loop agent'],
    ['', '', ''],
    ['Graph Core', '', ''],
    ['', 'spring-ai-alibaba-graph-core/.../graph/StateGraph.java', 'Workflow definition (677 lines)'],
    ['', 'spring-ai-alibaba-graph-core/.../graph/CompiledGraph.java', 'Execution engine (778 lines)'],
    ['', 'spring-ai-alibaba-graph-core/.../graph/OverAllState.java', 'State container (561 lines)'],
    ['', 'spring-ai-alibaba-graph-core/.../graph/Node.java', 'Graph node (130 lines)'],
    ['', 'spring-ai-alibaba-graph-core/.../graph/Edge.java', 'Graph edge (150 lines)'],
    ['', 'spring-ai-alibaba-graph-core/.../checkpoint/MemorySaver.java', 'In-memory checkpoint saver'],
    ['', '', ''],
    ['ManifoldCF', '', ''],
    ['', 'manifoldcf-saikat/connectors/cmis/connector/.../CmisRepositoryConnector.java', 'CMIS connector with ACL extraction'],
    ['', 'manifoldcf-saikat/connectors/cmis/connector/.../HttpsForceHttpInvoker.java', 'HTTPS URL rewriting'],
    ['', '', ''],
    ['Frontend', '', ''],
    ['', 'frontend/packages/main/src/pages/Source/Create/index.tsx', 'Source creation with vendor presets'],
    ['', 'frontend/packages/main/src/pages/Knowledge/Create/index.tsx', 'Knowledge base creation'],
    ['', 'frontend/packages/main/src/pages/Chat/index.tsx', 'Chat interface'],
    ['', 'frontend/packages/main/src/services/modelService.ts', 'Model API service layer'],
    ['', 'frontend/packages/main/src/layouts/SideMenuLayout.tsx', 'Sidebar navigation'],
]

add_table(['Category', 'File Path', 'Description'], [r for r in files if r[1]])

# ═══════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), 'Spring_AI_Alibaba_Architecture.docx')
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
print(f"   Pages: ~50+ pages")
print(f"   Sections: 13 major sections")
print(f"   Code snippets: 40+")
print(f"   Tables: 20+")
print(f"   Architecture diagrams: 10+")
